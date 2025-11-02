import docx
import re
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

class DocumentProcessor:
    """Handles extracting text and tables from DOCX files."""
    
    @staticmethod
    def extract_text_and_tables_from_docx(file_path: str) -> str:
        """Extract text content and tables from DOCX file with better table handling"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # Better table extraction
            for table_idx, table in enumerate(doc.tables):
                content.append(f"\n--- TABLE {table_idx + 1} ---")
                
                # Extract table data
                table_rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_rows.append(row_data)
                
                # Add structured table information
                if table_rows:
                    # Add headers
                    headers = table_rows[0]
                    content.append(f"HEADERS: {' | '.join(headers)}")
                    
                    # Add row data
                    for row_idx, row_data in enumerate(table_rows[1:], 1):
                        content.append(f"ROW {row_idx}: {' | '.join(row_data)}")
                    
                    # Add table summary for better search
                    content.append(f"TABLE_SUMMARY: {len(table_rows)-1} rows, {len(headers)} columns")
                    
                    # Add individual cell content for keyword search
                    for row_idx, row_data in enumerate(table_rows):
                        for col_idx, cell_content in enumerate(row_data):
                            if cell_content and len(cell_content) > 3:
                                content.append(f"CELL[{row_idx},{col_idx}]: {cell_content}")
                
                content.append("--- END TABLE ---\n")
            
            return '\n'.join(content)
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {e}")

    @staticmethod
    def split_document_text(text_content: str) -> list[Document]:
        """Split document text into chunks for processing."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )

        # Detect if this chunk contains table data
        contains_table = "--- TABLE" in text_content

        # Create a LangChain Document first
        doc = Document(
            page_content=text_content,
            metadata={
                "source": "uploaded_file",
                "type": "table" if contains_table else "document",
            }
        )
        return text_splitter.split_documents([doc])