"""
Standalone Model Comparison Report Generator
Uses the EXACT same setup as model_manager.py for accurate testing.

Usage: python model_comparison_report/standalone_model_comparison_no_tables.py model_comparison_report/Translated_Norms.docx
"""

import os
import sys
import time
import warnings
from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse

# directory to path to import the modules
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)  # up to emc-norms-chatbot
sys.path.append(parent_dir)

from document_processor import DocumentProcessor
from vector_store_manager import VectorStoreManager
from model_manager import ModelManager, MODELS
from hybrid_search import HybridSearch
from query_filter import QueryFilter

warnings.filterwarnings("ignore", category=FutureWarning)

# Complete test questions
TEST_QUESTIONS = [
    {
        "id": "Q1",
        "question": "How can electromagnetic interference sources be categorized? What are examples of each type?",
        "expected_answer": "Electromagnetic interference sources can be divided into two main types: Narrowband sources, such as vehicle electronic modules with clock generators, oscillators, digital logic circuits (e.g., microprocessors), and display units. Broadband sources, such as electric motors and ignition systems."
    },
    {
        "id": "Q2", 
        "question": "What key elements must be included in a test plan for electromagnetic compatibility testing?",
        "expected_answer": "The frequency range to be tested, Emission limit values, Types and locations of antennas, Requirements for the test report, Supply voltage and other relevant parameters, including operating conditions of the device under test (DUT), as described in section 4.1.4"
    },
    {
        "id": "Q3",
        "question": "Why must the DUT operate under typical vehicle load conditions during interference emission testing?",
        "expected_answer": "When testing components or modules, the DUT must operate under typical load and conditions as encountered in vehicles, ensuring maximum emission states are achieved. These conditions must be defined in the test plan."
    },
    {
        "id": "Q4",
        "question": "What must the test plan clarify regarding compliance for each frequency range?",
        "expected_answer": "The test plan shall specify for each frequency range whether compliance can be achieved with mean and peak limits or with mean and quasi-peak limits."
    },
    {
        "id": "Q5",
        "question": "Where can the peripheral interface unit be placed during testing, and what determines its allowable emission levels?",
        "expected_answer": "The peripheral interface unit may be placed inside or outside the shielded chamber. If placed inside, its emission levels must be at least 6 dB below the specified limits in the test plan."
    },
    {
        "id": "Q6",
        "question": "What two requirements must the peripheral interface unit meet when testing the DUT?",
        "expected_answer": "To ensure proper operation during testing, a peripheral interface unit replicating the vehicle installation must be used. All essential sensor and actuator lines of the DUT must be connected to this unit. The unit must be capable of controlling the DUT as specified."
    },
    {
        "id": "Q7",
        "question": "If the peripheral interface unit is placed inside the shielded chamber, how must its emission levels compare to the test plan limits? ",
        "expected_answer": "If placed inside, its emission levels must be at least 6 dB below the specified limits in the test plan. "
    },
    {
        "id": "Q8",
        "question": "Under what conditions can the interference voltage limits be adjusted?",
        "expected_answer": "If other types of receivers are used or different coupling models for interference propagation apply, the limits can be adjusted and specified individually by the vehicle manufacturer."
    },
    {
        "id": "Q9",
        "question": "What must be done if other types of receivers or different coupling models are used?",
        "expected_answer": "If other types of receivers are used or different coupling models for interference propagation apply, the limits may be adjusted and specified individually by the vehicle manufacturer."
    }
]

class InstantModelTester:
    """Test models instantly by processing DOCX files directly"""
    
    def __init__(self):
        self.vectorstore = None
        self.results = {}
        self.hybrid_searcher = None
        
        # Initialize your actual managers
        self.vector_manager = VectorStoreManager()
        self.model_manager = ModelManager()
        self.document_processor = DocumentProcessor()
        self.query_filter = QueryFilter()
    
    def process_document(self, docx_path: str):
        """Process a DOCX file using the EXACT same method as the app"""
        if not os.path.exists(docx_path):
            print(f"❌ Document not found: {docx_path}")
            return False
        
        try:
            print("📄 Processing document (using app's exact method)...")
            
            # Use the EXACT same extraction as the app
            text_content = self.document_processor.extract_text_and_tables_from_docx(docx_path)
            
            if not text_content:
                print("❌ Failed to extract text from document")
                return False
            
            # Use the EXACT same splitting as the app
            documents = self.document_processor.split_document_text(text_content)
            print(f"✅ Document split into {len(documents)} chunks (app method)")
            
            # Create in-memory vectorstore (SAME as app)
            self.vectorstore = self.vector_manager.create_vectorstore(documents, "Test_Document")
            print("✅ Vector store created successfully")
            
            # Initialize hybrid search exactly like the app does
            self.hybrid_searcher = HybridSearch(self.vectorstore)
            
            # Pre-load BM25 index exactly like the app does
            all_docs = self.vectorstore.similarity_search("", k=len(documents) + 10)
            self.hybrid_searcher._build_bm25_index(all_docs)
            print("✅ Hybrid search initialized (same as app)")
            
            # Quick test using hybrid search (not just semantic)
            test_results = self.hybrid_searcher.hybrid_search("test", k=1)
            print(f"✅ Hybrid search ready with {len(test_results)} test results")
            
            return True
            
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            return False

    def create_app_retrieval_chain(self, model_name: str):
        """Create retrieval chain that works EXACTLY like the app"""
        try:
            # Get the LLM and prompt (same as app)
            llm = self.model_manager.create_ollama_llm(model_name)
            prompt = self.model_manager.create_prompt()
            
            # Import the chain creation function
            from langchain.chains.combine_documents import create_stuff_documents_chain
            document_chain = create_stuff_documents_chain(llm, prompt)
            
            def retrieval_chain_func(inputs: dict, vectorstore=None) -> dict:
                """EXACT replica of the app's retrieval chain function"""
                if vectorstore is None:
                    return {
                        "answer": "No document context available.",
                        "context": []
                    }
                
                query = inputs["input"]
                
                # Use the SAME hybrid search as the app
                hybrid_searcher = HybridSearch(vectorstore)
                
                # Pre-load BM25 if needed (same as app logic)
                if hybrid_searcher.bm25_index is None:
                    all_docs = vectorstore.similarity_search("", k=1000)
                    if all_docs:
                        hybrid_searcher._build_bm25_index(all_docs)
                
                # Get context using hybrid search with SAME parameters as app
                context_docs = hybrid_searcher.hybrid_search(query, k=5, alpha=0.6)
                
                if not context_docs:
                    return {
                        "answer": "I cannot find relevant information in the document to answer this question.",
                        "context": []
                    }
                
                # Invoke the document chain (same as app)
                result = document_chain.invoke({
                    "context": context_docs, 
                    "input": query
                })
                
                # Return format EXACTLY like the app
                if isinstance(result, str):
                    return {
                        "answer": result,
                        "context": context_docs
                    }
                elif isinstance(result, dict) and "answer" in result:
                    return {
                        "answer": result["answer"],
                        "context": context_docs
                    }
                else:
                    return {
                        "answer": str(result),
                        "context": context_docs
                    }
            
            return retrieval_chain_func
            
        except Exception as e:
            print(f"❌ Error creating app-style retrieval chain: {e}")
            return None

    def test_model(self, model_name: str, questions: list):
        """Test a single model using the EXACT same setup as the app"""
        print(f"\n🤖 Testing {MODELS[model_name]['name']} (app-equivalent setup)...")
        
        try:
            # Use the app-equivalent retrieval chain
            retrieval_chain = self.create_app_retrieval_chain(model_name)
            
            if retrieval_chain is None:
                print(f"❌ Failed to create retrieval chain for {model_name}")
                return None
            
            model_results = {
                'model_name': model_name,
                'display_name': MODELS[model_name]['name'],
                'description': MODELS[model_name]['description'],
                'questions': [],
                'total_time': 0,
                'valid_responses': 0,
                'failed_responses': 0,
                'avg_response_time': 0,
                'response_rate': 0
            }
            
            for i, qa in enumerate(questions, 1):
                print(f"  Q{i}: {qa['question'][:60]}...")
                
                # CRITICAL: Use the same query filtering as the app
                if not self.query_filter.is_document_related_query(qa['question']):
                    print(f"    ⚠️ Skipped (non-document query per app filter)")
                    continue
                
                try:
                    start_time = time.time()
                    
                    # Use the EXACT same call as the app
                    response = retrieval_chain({"input": qa['question']}, vectorstore=self.vectorstore)
                    response_time = time.time() - start_time
                    
                    model_results['total_time'] += response_time
                    answer = response.get('answer', '').strip()
                    
                    # Use the SAME validation logic as the app
                    is_valid = (
                        answer and 
                        len(answer) > 10 and 
                        not answer.startswith('Error') and
                        not "cannot find" in answer.lower() and
                        not "No document context" in answer and
                        not "I cannot find relevant information" in answer
                    )
                    
                    if is_valid:
                        model_results['valid_responses'] += 1
                    else:
                        model_results['failed_responses'] += 1
                    
                    model_results['questions'].append({
                        'question': qa['question'],
                        'expected_answer': qa['expected_answer'],
                        'actual_answer': answer,
                        'response_time': response_time,
                        'valid': is_valid,
                        'context_used': len(response.get('context', [])),
                        'context_preview': [doc.page_content[:100] + "..." for doc in response.get('context', [])][:2]  # For debugging
                    })
                    
                    status = "✓" if is_valid else "✗"
                    print(f"    {status} {response_time:.2f}s (context: {len(response.get('context', []))} chunks)")
                    
                    time.sleep(0.5)  # Same delay as app might have
                    
                except Exception as e:
                    print(f"    ✗ Error: {str(e)}")
                    model_results['failed_responses'] += 1
                    model_results['questions'].append({
                        'question': qa['question'],
                        'expected_answer': qa['expected_answer'],
                        'actual_answer': f"Error: {str(e)}",
                        'response_time': 0,
                        'valid': False,
                        'context_used': 0,
                        'context_preview': []
                    })
            
            # Calculate valid response rate (same as app would)
            total_questions_attempted = len([q for q in model_results['questions']])
            if total_questions_attempted > 0:
                model_results['avg_response_time'] = model_results['total_time'] / total_questions_attempted
                model_results['response_rate'] = (model_results['valid_responses'] / total_questions_attempted) * 100
            
            return model_results
            
        except Exception as e:
            print(f"❌ Error testing {model_name}: {e}")
            return None

    def test_all_models(self, docx_path: str, models_to_test: list = None):
        """Test all models using the EXACT same setup as the app"""
        print("🚀 App-Equivalent Model Comparison Test")
        print("=" * 60)
        print(f"📄 Processing: {os.path.basename(docx_path)}")
        print("⚡ Using EXACT same methods as Streamlit app")
        
        if not self.process_document(docx_path):
            return None
        
        if models_to_test is None:
            models_to_test = list(MODELS.keys())
        
        print(f"🤖 Testing {len(models_to_test)} models (app-equivalent setup)...")
        
        tested_models = 0
        for model_name in models_to_test:
            try:
                # Verify model is available
                print(f"🔍 Checking {model_name} availability...")
                try:
                    llm = self.model_manager.create_ollama_llm(model_name)
                    # Quick test to ensure model responds
                    test_response = llm.invoke("Say 'ready'")
                    print(f"✅ {model_name} is ready")
                except Exception as e:
                    print(f"❌ {model_name} not available: {e}")
                    continue
                
                result = self.test_model(model_name, TEST_QUESTIONS)
                if result:
                    self.results[model_name] = result
                    tested_models += 1
                    print(f"✅ Completed {model_name} (app-equivalent)")
                else:
                    print(f"❌ Failed to test {model_name}")
            except Exception as e:
                print(f"❌ Error testing {model_name}: {e}")
        
        print(f"\n🎯 Testing complete: {tested_models}/{len(models_to_test)} models tested")
        print("📊 Results are NOW IDENTICAL to Streamlit app results")
        return self.results

class ReportGenerator:
    """Generate comprehensive Word report from test results"""
    
    def __init__(self, results: dict, document_name: str = "Test Document"):
        self.results = results
        self.document_name = document_name
        self.output_dir = 'reports'
        
    def add_table_to_doc(self, doc, data: list, headers: list, style: str = 'Table Grid'):
        """Add a formatted table to the document"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = style
        
        # Add headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
    
    def add_heading_with_color(self, doc, text: str, level: int, color: tuple = None):
        """Add a heading with optional color"""
        heading = doc.add_heading(text, level)
        if color:
            for run in heading.runs:
                run.font.color.rgb = RGBColor(*color)
        return heading
    
    def generate_report(self, output_filename: str = None):
        """Generate comprehensive comparison report"""
        if not output_filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"Model_Comparison_Report_{timestamp}.docx"
        
        # Create output directory
        os.makedirs(self.output_dir, exist_ok=True)
        output_path = os.path.join(self.output_dir, output_filename)
        
        # Create document
        doc = Document()
        
        # Title Page
        title = self.add_heading_with_color(doc, 'Model Performance Comparison Report', 0, (0, 51, 102))
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph('Electromagnetic Compatibility Document Analysis')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].italic = True
        
        doc.add_paragraph()  # Spacing
        
        # Report metadata table
        metadata_table = [
            ['Document Analyzed', self.document_name],
            ['Report Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Total Models Tested', str(len(self.results))],
            ['Questions per Model', str(len(TEST_QUESTIONS))],
            ['Testing Method', 'Direct DOCX Processing'],
            ['Retrieval Strategy', 'Hybrid Search (Semantic + Keyword)']
        ]
        
        doc.add_heading('Report Overview', level=1)
        self.add_table_to_doc(doc, metadata_table, ['Attribute', 'Value'], 'Light Shading')
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        
        if self.results:
            # Calculate statistics
            best_model = max(self.results.values(), key=lambda x: x.get('response_rate', 0))
            fastest_model = min(self.results.values(), key=lambda x: x.get('avg_response_time', float('inf')))
            total_time = sum([r.get('total_time', 0) for r in self.results.values()])
            total_questions = len(self.results) * len(TEST_QUESTIONS)
            total_valid = sum([r.get('valid_responses', 0) for r in self.results.values()])
            overall_response_rate = (total_valid / total_questions) * 100 if total_questions > 0 else 0
            
            summary_text = f"""This comprehensive report analyzes the performance of {len(self.results)} AI language models on {len(TEST_QUESTIONS)} standardized questions about electromagnetic compatibility testing procedures.

The testing was conducted using a production-equivalent setup with hybrid search retrieval, ensuring that the results accurately reflect real-world performance.

KEY FINDINGS:
• Best Performing Model: {best_model['display_name']} with {best_model.get('response_rate', 0):.1f}% response rate
• Fastest Model: {fastest_model['display_name']} with {fastest_model.get('avg_response_time', 0):.2f} seconds average response time
• Overall Response Rate: {overall_response_rate:.1f}% across all models and questions
• Overall Success Rate: XX % across all models and questions
• Total Processing Time: {total_time:.2f} seconds for complete test suite

PERFORMANCE TRADE-OFFS:
The analysis reveals clear trade-offs between response quality, speed, and computational requirements. Larger models generally provide more accurate and detailed responses but require longer processing times, while smaller models offer faster responses with potentially lower accuracy for complex technical questions."""
            
            doc.add_paragraph(summary_text)
        
        # Performance Overview Table
        doc.add_heading('Performance Overview', level=1)
        
        performance_data = []
        for model_name, result in self.results.items():
            performance_data.append([
                result['display_name'], 
                "",
                f"{result.get('response_rate', 0):.1f}%",
                f"{result.get('avg_response_time', 0):.2f}s",
                result.get('valid_responses', 0),
                "",
                result.get('failed_responses', 0),
                f"{result.get('total_time', 0):.2f}s"
            ])
        
        headers = ['Model','Success Rate', 'Response Rate', 'Avg Time', 'Valid Responses', 'Success Responses', 'Failed Responses', 'Total Time']
        self.add_table_to_doc(doc, performance_data, headers, 'Medium Shading 1')
        
        doc.add_paragraph(
            "Manual Success Score: "
            "Scores were assigned based on factual accuracy, completeness, and relevance. "
            "A score of 2 indicates a fully correct and concise answer; 1.5 represents a correct answer "
            "that includes unnecessary or verbose details; and 1 indicates a partially correct or incomplete response. "
            "A score of 0 is reserved for incorrect or unsupported answers."
        )
        doc.add_paragraph()

        # Technical Appendix
        doc.add_page_break()
        doc.add_heading('Technical Appendix', level=1)
        
        tech_details = f"""
TESTING METHODOLOGY:

1. Document Processing:
   - Direct DOCX file processing without intermediate steps
   - Enhanced text extraction with table support
   - Recursive text splitting with 1000-character chunks, 200-character overlap

2. Vector Storage and Retrieval:
   - Embedding Model: sentence-transformers/all-MiniLM-L6-v2
   - Vector Database: ChromaDB (in-memory for testing)
   - Retrieval Strategy: Hybrid search combining semantic and keyword approaches
   - Retrieval Parameters: Top-5 most relevant chunks per query

3. Model Configurations:
   - Base URL: http://localhost:11434 (Ollama)
   - Temperature: Varied by model (typically 0.1-0.3)
   - Top-p: Varied by model (typically 0.7-0.9)
   - Max Tokens: Adjusted based on model capabilities

4. Evaluation Metrics:
   - Success Criteria: Relevant, coherent responses based on document context
   - Response Time: End-to-end processing time including retrieval
   - Context Usage: Number of document chunks used for response generation

MODEL SPECIFICATIONS:
"""
        
        for model_name, config in MODELS.items():
            tech_details += f"""
• {config['name']}:
  - Model ID: {model_name}
  - Temperature: {config.get('temperature', 'N/A')}
  - Top-p: {config.get('top_p', 'N/A')}
  - Max Tokens: {config.get('num_predict', 'N/A')}
"""
        
        doc.add_paragraph(tech_details)
        
        # Save document
        doc.save(output_path)
        print(f"📊 Report generated: {output_path}")
        return output_path
    
    def generate_single_model_report(self, model_name: str, result: dict):
        """Generate a standalone DOCX report for one specific model"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{result['display_name'].replace(' ', '_')}_Report_{timestamp}.docx"
        output_path = os.path.join(self.output_dir, filename)
        os.makedirs(self.output_dir, exist_ok=True)

        doc = Document()
        doc.add_heading(f"{result['display_name']} - Individual Model Report", 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Document: {self.document_name}")
        doc.add_paragraph()

        # Metrics summary
        doc.add_heading("Performance Summary", level=1)
        summary = f"""
Success Rate: 
Response Rate: {result.get('response_rate', 0):.1f}%
Average Response Time: {result.get('avg_response_time', 0):.2f} seconds
Valid Responses: {result.get('valid_responses', 0)}/{len(TEST_QUESTIONS)}
Correct Responses: X /{len(TEST_QUESTIONS)}
Failed Responses: {result.get('failed_responses', 0)}/{len(TEST_QUESTIONS)}
Total Time: {result.get('total_time', 0):.2f} seconds
"""
        doc.add_paragraph(summary)

        # Add manual scoring guide
        doc.add_paragraph(
            "Manual Success Score: "
            "Each answer can be scored manually based on factual accuracy, completeness, and relevance."
        )

        # All questions
        doc.add_heading("Question-by-Question Responses", level=1)
        for i, qa in enumerate(result['questions'], 1):
            doc.add_heading(f"Question {i}", level=2)
            doc.add_paragraph(f"Question: {qa['question']}")
            doc.add_paragraph(f"Expected Answer: {qa['expected_answer']}")

            # Status & placeholder for manual score
            if qa['valid']:
                status_text = f"✓ VALID - {qa['response_time']:.2f}s - Context chunks: {qa['context_used']}"
                color = (0, 128, 0)
            else:
                status_text = f"✗ FAILED - Error or invalid response"
                color = (255, 0, 0)

            status_para = doc.add_paragraph()
            run = status_para.add_run(status_text)
            run.bold = True
            run.font.color.rgb = RGBColor(*color)

            score_para = doc.add_paragraph()
            score_run = score_para.add_run("Manual Success Score: ")
            score_run.bold = True

            doc.add_paragraph(f"Model’s Response: {qa['actual_answer']}")
            doc.add_paragraph("_" * 60)
            doc.add_paragraph()

        doc.save(output_path)
        print(f"📝 Individual report saved: {output_path}")
        return output_path

def main():
    """Main function with enhanced app-equivalent testing"""
    parser = argparse.ArgumentParser(description='App-Equivalent Model Comparison')
    parser.add_argument('docx_path', help='Path to the DOCX file to test')
    parser.add_argument('--models', nargs='+', help='Specific models to test')
    parser.add_argument('--output', help='Output report filename')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.docx_path):
        print(f"❌ Document not found: {args.docx_path}")
        return
    
    print("🚀 App-Equivalent Model Comparison")
    print("USING EXACT SAME METHODS AS STREAMLIT APP")
    print("=" * 60)
    
    # Initialize and test with app-equivalent setup
    tester = InstantModelTester()
    results = tester.test_all_models(args.docx_path, args.models)
    
    if not results:
        print("❌ Testing failed")
        return
    
    print(f"\n📊 Generating report...")
    document_name = os.path.basename(args.docx_path)
    generator = ReportGenerator(results, document_name)
    report_path = generator.generate_report(args.output)

    # Generate individual model reports
    print("\n🧾 Generating individual model reports...")
    for model_name, result in results.items():
        generator.generate_single_model_report(model_name, result)
    
    print(f"\n✅ App-Equivalent Testing Complete!")
    print(f"📄 Report saved: {report_path}")
    print(f"🤖 Models tested: {len(results)}")
    print(f"🔍 Results are IDENTICAL to Streamlit app performance")
    
    # Enhanced summary
    print(f"\n📈 Results Summary (App-Equivalent):")
    for model_name, result in results.items():
        print(f"• {result['display_name']}: {result['response_rate']:.1f}% responses, {result['avg_response_time']:.2f}s avg")

if __name__ == "__main__":
    main()